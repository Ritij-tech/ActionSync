import subprocess
import sys
import json
import os
from datetime import date
from docx import Document


def run(script: str) -> None:
    subprocess.run([sys.executable, script], check=True)


def replace_placeholders(doc: Document, mapping: dict) -> None:
    for paragraph in doc.paragraphs:
        for key, value in mapping.items():
            placeholder = f"{{{key}}}"
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in mapping.items():
                        placeholder = f"{{{key}}}"
                        if placeholder in paragraph.text:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))


if __name__ == "__main__":
    print("Running Scraper.py...")
    run("Scraper.py")

    print("Running LLM_output.py...")
    run("LLM_output.py")

    with open("openai_output.json", "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document("Digital_Presence_Audit_Template.docx")

    page1 = {
        "ProspectName": data.get("ProspectName", "Prospect"),
        "AuditDate": date.today().strftime("%Y-%m-%d"),
    }
    replace_placeholders(doc, page1)

    page2 = {
        "StrategicInsight": data.get("StrategicInsight", ""),
        "CXOPositioning": data.get("CXOPositioning", ""),
    }
    replace_placeholders(doc, page2)

    page3 = {
        "ProfilePhotoStatus": data.get("ProfilePhotoStatus", ""),
        "ProfilePhotoFixable": data.get("ProfilePhotoFixable", ""),
        "HeadlineStatus": data.get("HeadlineStatus", ""),
        "HeadlineFixable": data.get("HeadlineFixable", ""),
        "CoverImageStatus": data.get("CoverImageStatus", ""),
        "CoverImageFixable": data.get("CoverImageFixable", ""),
        "AboutNarrativeStatus": data.get("AboutNarrativeStatus", ""),
        "AboutFixable": data.get("AboutFixable", ""),
        "FirstImpressionSummary": data.get("FirstImpressionSummary", ""),
    }
    replace_placeholders(doc, page3)

    page4 = {
        "TotalFollowers": data.get("TotalFollowers", ""),
        "PeerFollowers": data.get("PeerFollowers", ""),
        "PostsPerMonth": data.get("PostsPerMonth", ""),
        "EngagementRate": data.get("EngagementRate", ""),
        "DaysSinceLastPost": data.get("DaysSinceLastPost", ""),
    }
    replace_placeholders(doc, page4)

    page5 = {
        "AuthorityScore": data.get("AuthorityScore", ""),
        "AuthorityNotes": data.get("AuthorityNotes", ""),
        "ProofScore": data.get("ProofScore", ""),
        "ProofNotes": data.get("ProofNotes", ""),
        "FounderVisibilityScore": data.get("FounderVisibilityScore", ""),
        "FounderNotes": data.get("FounderNotes", ""),
        "TrustHookScore": data.get("TrustHookScore", ""),
        "TrustHookNotes": data.get("TrustHookNotes", ""),
    }
    replace_placeholders(doc, page5)

    page6 = {
        "FixableGap1": data.get("FixableGap1", ""),
        "FixableGap2": data.get("FixableGap2", ""),
        "FixableGap3": data.get("FixableGap3", ""),
    }
    replace_placeholders(doc, page6)

    os.makedirs("output", exist_ok=True)
    filename = os.path.join(
        "output",
        f"Presence_Audit_{data.get('ProspectName', 'Prospect').replace(' ', '_')}.docx",
    )
    doc.save(filename)

    print(f"Presence Audit saved as {filename}")
